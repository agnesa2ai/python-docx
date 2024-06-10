# cloned git repo from: https://github.com/python-openxml/python-docx?tab=readme-ov-file



#run in terminal: pip install python-docx
import os
import re
import docx
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK


#bring in shell 
document = Document('python-docx/pre.docx')

#establish a pattern to magic string ..///
start_pattern = r'\.\.///'
end_pattern = r'\.png'
magic_pattern = re.compile(start_pattern + '.*?' + end_pattern)


#locate magic string, insert image, insert source path
document.paragraphs  # to extract paragraphs
for par in document.paragraphs: 
    matches = re.findall(start_pattern, par.text)
    if matches:
        for match in matches:
            match_index = par.text.find(match) #locates a magic string based on start_pattern
            figure_name = par.text[match_index + len(match):] #extracts the figure name from the magic string
            figure_location_and_name = ("python-docx/images/" + figure_name)  #can change "images" later. but don't know if necessary
            magic_matches = magic_pattern.findall(par.text) #relocates magic string this time from start_pattern to end_pattern 
            for match in magic_matches:
                par.text = par.text.replace(match, "")
            r = par.add_run('')  
            r.add_picture(figure_location_and_name, width=Inches(2)) #inserting image 
            rr = par.add_run('\n')
            rr.add_text('Source: ' + figure_location_and_name) #inserting source path as text
            
#render out draft
document.save('post.docx')

